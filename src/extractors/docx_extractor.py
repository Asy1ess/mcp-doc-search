"""Word DOCX text extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import Table

from src.extractors._utils import make_document


def _table_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            rows.append("\t".join(cells))
    return "\n".join(rows)


def extract_docx(path: Path) -> ExtractedDocument:
    document = Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        table_body = _table_text(table)
        if table_body:
            parts.append(table_body)

    return make_document(path, "\n\n".join(parts), format="docx")
